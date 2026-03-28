from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# --- Styles ---
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)


# Helper: add a shaded heading row to a table cell
def shade_cell(cell, hex_color="1F3864"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


# ---- TITLE ----
title = doc.add_heading("Data Management Plan", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

subtitle = doc.add_paragraph(
    "Longitudinal Assessment of the Relationship Between AI-Assisted Coding Practices and Python Coding Skill",
)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(12)

doc.add_paragraph()

# ---- PROJECT DETAILS TABLE ----
details = [
    ("Creator", "Andy Woods"),
    ("Principal Investigator", "Dr Andy Woods"),
    ("Data Manager", "Dr Andy Woods"),
    ("Project Team", "Dr Andy Woods, Dr Alex Reppel, Chris Chowen"),
    ("Affiliation", "Royal Holloway, University of London"),
    ("Template", "Royal Holloway Generic DMP"),
    ("Grant Number", "N/A (unfunded internal project)"),
    ("Last Modified", "21-03-2026"),
]

tbl = doc.add_table(rows=len(details), cols=2)
tbl.style = "Table Grid"
tbl.columns[0].width = Inches(2.2)
tbl.columns[1].width = Inches(4.3)

for i, (label, value) in enumerate(details):
    row = tbl.rows[i]
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = value
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    for cell in row.cells:
        set_cell_border(cell)

doc.add_paragraph()

# ---- PROJECT ABSTRACT ----
h = doc.add_heading("Project Abstract", level=2)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

abstract = (
    'This study investigates whether reliance on AI-assisted coding tools ("vibe coding") leads to measurable '
    "changes in Python coding skill over time. Participants who already code in Python periodically complete short "
    "Python coding challenges on a purpose-built web platform (agenticbrainrot.com) and self-report their AI coding "
    "habits via a brief post-session survey. Sessions are separated by a minimum of 28 days over a 12-month data "
    "collection period. A qualitative strand (exit survey and optional follow-up interviews) explores participants' "
    "own explanations for any skill changes. The study follows a citizen-science model: participants can view their "
    "own performance trends, and the fully anonymised dataset will be released publicly after a 12-month embargo. "
    "The study is pre-registered on the Open Science Framework (OSF) before data collection begins."
)
doc.add_paragraph(abstract)

doc.add_paragraph()

# ---- DMP SECTIONS ----
sections_data = [
    (
        "Data Summary",
        "1. Briefly introduce the types of data the research will create. Why did you decide to use these data types?",
        "The study will generate three types of data.\n\n"
        "First, structured quantitative performance data: for each coding challenge attempt, the platform automatically records the proportion of test cases passed (accuracy), active completion time in seconds, idle time, device type, and technical issue flags. This approach was chosen because automatic capture eliminates transcription error and produces a precise, objective measure of coding skill that can be tracked longitudinally.\n\n"
        "Second, structured survey data: participants self-report their AI coding habits (percentage of coding time that is AI-assisted, how thoroughly they review AI-generated code), hours coding per week, and demographic/background information. Surveys are used because the primary independent variable  -  vibe-coding behaviour  -  cannot be observed directly and must be self-reported.\n\n"
        "Third, qualitative free-text data: an optional exit survey (five open questions) and optional semi-structured interviews (~15–20 participants) explore participants' own explanations for any skill changes. These are included to contextualise and interrogate the quantitative findings using a convergent mixed-methods design.",
    ),
    (
        "Data Collection",
        "2. Give details on the proposed methodologies that will be used to create the data.",
        "Performance data is captured automatically by the platform: Python code is executed client-side in the browser using Pyodide (WebAssembly), and results (pass/fail per test case, timestamps) are sent to the Django/PostgreSQL backend over HTTPS. No code is executed on the server. Survey responses are collected via in-app forms immediately after each session. Consent is obtained digitally before first participation and is versioned and timestamped.\n\n"
        "To ensure data quality: attempts flagged with technical_issues=True are retained in the database but excluded from the primary analysis. Sessions with fewer than 4 challenge attempts are retained but excluded from the multilevel model. Implausible timing values (active_time_seconds < 10 or > 3600) are excluded from the completion-speed analysis. The study is pre-registered on OSF before data collection begins; any deviations are logged in a dedicated deviations file in the project repository and reported in the published paper.\n\n"
        "Optional interviews will be conducted via Zoom or equivalent, recorded with explicit participant consent (separate from study consent), and transcribed either by AI with human correction or entirely by a human researcher. Transcripts will be coded using reflexive thematic analysis (Braun & Clarke), supported by appropriate tools (e.g. NVivo, MaxQDA). The team has extensive experience in both qualitative methods and secure web application development.",
    ),
    (
        "Short-term Data Storage",
        "3. How will the data be stored in the short term?",
        "Performance and survey data are stored in a PostgreSQL database hosted on a Hetzner VPS (Virtual Private Server) in the EU (Germany). Hetzner operates under a data processing agreement with Royal Holloway and meets GDPR requirements. The platform is built using Django, following established security best practices. Access to the production database is restricted to the researcher (Dr Andy Woods) and, where necessary, authorised technical support. Data is separated by design: identifiable account data (email addresses) is held in a separate table from pseudonymous research data (participant ID only). All data in transit is encrypted via HTTPS (TLS). Interview recordings are held in an encrypted folder on the researcher's university-managed device and are not stored in the cloud. Storage is in line with institutional data management policy.",
    ),
    (
        "Short-term Data Storage",
        "3a. What backup will you have in the in-project period to ensure no data is lost?",
        "The PostgreSQL database is automatically backed up daily via the Appliku deployment platform, with snapshots retained for 30 days. At regular intervals during the study (at minimum every 3 months), an anonymised export of the research data will be downloaded and stored in Royal Holloway OneDrive (Microsoft Azure Cloud, operated under RHUL's service agreement with Microsoft, meeting GDPR requirements). Interview recordings will be backed up to an encrypted RHUL OneDrive folder immediately after each recording, prior to transcription.",
    ),
    (
        "Long-term Data Storage",
        "4. How will the data be stored in the long term?",
        "On completion of the study, pseudonymous research data will be exported from the production database, fully anonymised (email addresses and any identifiable free-text content removed or redacted), and archived. The anonymised dataset will be deposited in the Open Science Framework (OSF) repository under the project's pre-registration DOI, and/or Zenodo, after a 12-month embargo from the date of first data collection. Analysis scripts (R) will be archived to GitHub (public) and Zenodo at the time of paper submission. Interview recordings will be deleted within 6 months of the research taking place, after transcription is verified. Pseudonymised interview transcripts will be stored in RHUL OneDrive.",
    ),
    (
        "Long-term Data Storage",
        "4a. Where have you decided to store it, why is this appropriate?",
        "The anonymised open dataset will be stored on OSF and/or Zenodo  -  both are established open-access repositories used by the research community for data sharing in the social and computer sciences, and both provide persistent DOIs. Analysis code will be stored on GitHub (public repository) and archived to Zenodo to ensure long-term reproducibility. Identifiable data and interview transcripts will be stored in RHUL OneDrive, in line with institutional guidelines.",
    ),
    (
        "Long-term Data Storage",
        "4b. How long will it be stored for and why?",
        "The anonymised open dataset and analysis scripts will be stored for a minimum of 10 years after publication, in line with RHUL data retention policy and open science norms. Interview recordings will be deleted within 6 months of the research taking place. Pseudonymised interview transcripts will be retained for 10 years after publication. Identifiable account data (email addresses) will be deleted within 12 months of the participant's withdrawal or the end of the study, whichever comes first.",
    ),
    (
        "Long-term Data Storage",
        "4c. Costs of storage  -  why are these appropriate?",
        "Hetzner VPS hosting costs approximately €10–20/month during active data collection (~24 months); this is a nominal cost borne by the researcher. OSF, Zenodo, and GitHub storage are free. RHUL OneDrive is provided by the institution at no additional cost. No significant storage budget is required.",
    ),
    (
        "Data Sharing",
        "5. How will the data be shared and the value it will have to others?",
        "A fully anonymised version of the dataset will be released publicly after a 12-month embargo under an open licence (CC BY 4.0 or equivalent), via OSF and/or Zenodo. This citizen-science dataset  -  the first longitudinal dataset linking self-reported AI coding habits to objective coding performance  -  will be of value to researchers in human-computer interaction, cognitive psychology, software engineering education, and AI tool design. Analysis scripts will be released on GitHub under MIT licence to support reproducibility and independent reanalysis. Summary findings will be shared via academic publication, conference presentations, and the study website.",
    ),
    (
        "Data Sharing",
        "5a. How will the data enhance the area and how could it be used in the future?",
        "The open dataset will enable independent reanalysis and meta-analytic work on the relationship between AI tool use and human skill. Future researchers could extend the analysis to other programming languages, different skill domains, or longer time horizons. The challenge content (Python coding problems with automated test suites) could also be reused by researchers studying coding skill assessment more broadly.",
    ),
    (
        "Data Sharing",
        "5b. When will you release the data?",
        "The anonymised dataset will be released after a 12-month embargo from the date of the first completed participant session. This embargo allows time for the primary analysis and publication before open release, while still meeting open science expectations. The pre-registered analysis plan is publicly available on OSF from the start of data collection. Results will be disseminated via academic journal publication and open-access preprint (e.g. PsyArXiv or arXiv) within 6 months of primary analysis.",
    ),
    (
        "Data Sharing",
        "5c. Will the data need to be updated?",
        "Participants may withdraw consent at any time before dataset anonymisation  -  in this case, their data will be deleted within 10 working days. After the dataset is anonymised and released, individual records can no longer be identified or removed. This is clearly explained in the participant information sheet and consent form.",
    ),
    (
        "Data Sharing",
        "5d. Will the data be open or will you charge for it?",
        "The anonymised dataset will be open and free to access under CC BY 4.0. Interview transcripts will not be released publicly.",
    ),
    (
        "Data Sharing",
        "5e. Financial requirements of sharing?",
        "Data will be shared via OSF, Zenodo, and GitHub, all of which are free to use. No financial implications.",
    ),
    (
        "Ethical and Legal Considerations",
        "6a. Any legal and ethical considerations of collecting the data?",
        "Ethical approval is being sought from the Royal Holloway Research Ethics Committee before data collection begins. The study is subject to UK GDPR; the lawful basis for processing personal data is consent. Participants' email addresses are the only identifiable data held and are stored separately from pseudonymous research data. No special category data (as defined by GDPR) is collected. Participants must give explicit digital informed consent before their first session; consent is versioned and timestamped. The study is pre-registered on OSF to ensure transparency and research integrity. IP in the platform and analysis code is held by Royal Holloway.",
    ),
    (
        "Ethical and Legal Considerations",
        "6b. Legal and ethical considerations around releasing and storing the data?",
        "Data collection, storage, and dissemination conform to Royal Holloway's ethical and data management policies. Participants receive an information sheet and consent form before participation, outlining their rights and the team's obligations. The three most important promises to participants are: (a) email addresses and identifiable account data will be deleted within 12 months of study end or withdrawal; (b) all data relating to a participant will be deleted within 10 working days if they withdraw, provided this is before dataset anonymisation; and (c) interview recordings will be deleted within 6 months of the research taking place, after transcription is verified. Free-text exit survey responses will be reviewed for inadvertent identifiability before public release. Interview transcripts will never be included in the public dataset.",
    ),
]

# Group by section

current_section = None
for section_name, question, answer in sections_data:
    if section_name != current_section:
        current_section = section_name
        h = doc.add_heading(section_name, level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        h.runs[0].font.size = Pt(13)

    # Question in bold
    q_para = doc.add_paragraph()
    q_run = q_para.add_run(question)
    q_run.bold = True
    q_run.font.size = Pt(11)
    q_run.font.color.rgb = RGBColor(0x2E, 0x4D, 0x7B)

    # Answer  -  handle embedded newlines
    for i, chunk in enumerate(answer.split("\n\n")):
        a_para = doc.add_paragraph(chunk.strip())
        a_para.paragraph_format.space_after = Pt(4)
        a_para.runs[0].font.size = Pt(11)

    doc.add_paragraph()

# ---- Save ----
out_path = "/Users/andytwoods/PycharmProjects/agenticbrainrot/ethics/data-management-plan-FILLED.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
